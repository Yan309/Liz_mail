"""CV parsing utilities to extract emails from PDF, Word, and ZIP files."""
import re
import os
import zipfile
import tempfile
from typing import List, Set
import PyPDF2
import pdfplumber
from docx import Document
from pathlib import Path
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class CVParser:
    """Extract email addresses from various document formats."""
    
    # Comprehensive email regex pattern
    EMAIL_PATTERN = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    
    @staticmethod
    def extract_emails_from_text(text: str) -> Set[str]:
        """Extract email addresses from text using regex."""
        emails = set(re.findall(CVParser.EMAIL_PATTERN, text))
        # Filter out common false positives
        valid_emails = {
            email.lower() for email in emails 
            if not any(exclude in email.lower() for exclude in ['example.com', 'test.com', 'domain.com'])
        }
        return valid_emails
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Set[str]:
        """Extract emails from PDF files using multiple methods."""
        emails = set()
        
        # Method 1: PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        emails.update(CVParser.extract_emails_from_text(text))
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {e}")
        
        # Method 2: pdfplumber (more accurate for complex PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        emails.update(CVParser.extract_emails_from_text(text))
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed for {file_path}: {e}")
        
        logger.info(f"Extracted {len(emails)} emails from PDF: {file_path}")
        return emails
    
    @staticmethod
    def extract_from_word(file_path: str) -> Set[str]:
        """Extract emails from Word documents (.doc, .docx)."""
        emails = set()
        
        try:
            doc = Document(file_path)
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                emails.update(CVParser.extract_emails_from_text(paragraph.text))
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        emails.update(CVParser.extract_emails_from_text(cell.text))
            
            logger.info(f"Extracted {len(emails)} emails from Word: {file_path}")
        except Exception as e:
            logger.error(f"Error extracting from Word document {file_path}: {e}")
        
        return emails
    
    @staticmethod
    def extract_from_zip(file_path: str) -> Set[str]:
        """Extract emails from all supported files within a ZIP archive."""
        emails = set()
        
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                # Create temporary directory for extraction
                with tempfile.TemporaryDirectory() as temp_dir:
                    zip_ref.extractall(temp_dir)
                    
                    # Process all files in the extracted directory
                    for root, _, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            file_ext = Path(file).suffix.lower()
                            
                            try:
                                if file_ext == '.pdf':
                                    emails.update(CVParser.extract_from_pdf(file_path))
                                elif file_ext in ['.doc', '.docx']:
                                    emails.update(CVParser.extract_from_word(file_path))
                                elif file_ext == '.txt':
                                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                        emails.update(CVParser.extract_emails_from_text(f.read()))
                            except Exception as e:
                                logger.warning(f"Failed to process {file} in ZIP: {e}")
            
            logger.info(f"Extracted {len(emails)} emails from ZIP: {file_path}")
        except Exception as e:
            logger.error(f"Error extracting from ZIP file {file_path}: {e}")
        
        return emails
    
    @staticmethod
    def process_file(file_path: str) -> Set[str]:
        """Process a single file and extract emails based on file type."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return CVParser.extract_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return CVParser.extract_from_word(file_path)
        elif file_ext == '.zip':
            return CVParser.extract_from_zip(file_path)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return CVParser.extract_emails_from_text(f.read())
        else:
            logger.warning(f"Unsupported file type: {file_ext}")
            return set()
    
    @staticmethod
    def process_multiple_files(file_paths: List[str]) -> List[dict]:
        """
        Process multiple files and return structured results.
        
        Returns:
            List of dicts with 'file', 'emails', and 'count' keys
        """
        results = []
        all_emails = set()
        
        for file_path in file_paths:
            try:
                emails = CVParser.process_file(file_path)
                all_emails.update(emails)
                results.append({
                    'file': Path(file_path).name,
                    'emails': list(emails),
                    'count': len(emails)
                })
            except Exception as e:
                logger.error(f"Error processing {file_path}: {e}")
                results.append({
                    'file': Path(file_path).name,
                    'emails': [],
                    'count': 0,
                    'error': str(e)
                })
        
        return results, list(all_emails)
